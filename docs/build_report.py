# -*- coding: utf-8 -*-
"""Generate the project report (.docx) for My Wallet - Personal Expense Tracker."""
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AMBER = RGBColor(0xF5, 0xA6, 0x00)
DARK = RGBColor(0x22, 0x22, 0x22)
GREY = RGBColor(0x66, 0x66, 0x66)
LIGHT_SHADE = "F2F2F2"
HEAD_SHADE = "FFF3CD"
CODE_SHADE = "F6F6F6"

doc = Document()

# ---------- Base styles ----------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for lvl, sz in [("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
    st = doc.styles[lvl]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.color.rgb = DARK
    st.font.bold = True

# ---------- Page setup (A4) ----------
sec = doc.sections[0]
sec.page_width = Cm(21.0)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.4)
sec.bottom_margin = Cm(2.4)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)

CONTENT_WIDTH = Cm(21.0 - 5.0)  # usable width


# ---------- Helpers ----------
def set_shade(element, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    element.append(shd)


def shade_cell(cell, hex_color):
    set_shade(cell._tc.get_or_add_tcPr(), hex_color)


def set_cell_borders(cell, color="CCCCCC", sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), sz)
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)
    tcPr.append(borders)


def para_border(p, color="DDDDDD", sz="6"):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), sz)
        e.set(qn("w:space"), "6")
        e.set(qn("w:color"), color)
        pbdr.append(e)
    pPr.append(pbdr)


def add_para(text="", size=11, bold=False, italic=False, color=None,
             align=None, space_after=6, space_before=0, style=None):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        if color:
            r.font.color.rgb = color
    return p


def add_bullets(items, style="List Bullet"):
    for it in items:
        p = doc.add_paragraph(style=style)
        if isinstance(it, tuple):
            r = p.add_run(it[0])
            r.font.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)
        p.paragraph_format.space_after = Pt(3)


def add_numbers(items):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(it)
        p.paragraph_format.space_after = Pt(3)


def add_code(code, caption=None):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = CONTENT_WIDTH
    shade_cell(cell, CODE_SHADE)
    set_cell_borders(cell, color="E0E0E0", sz="4")
    first = True
    for line in code.split("\n"):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    if caption:
        add_para(caption, size=9, italic=True, color=GREY,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10, space_before=2)
    else:
        add_para("", space_after=4)


def add_placeholder(label, caption, width_cm=8.0, height_cm=5.2):
    """A bordered grey box that stands in for a screenshot."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Cm(width_cm)
    shade_cell(cell, "FAFAFA")
    set_cell_borders(cell, color="BBBBBB", sz="8")
    # vertical centering
    tcPr = cell._tc.get_or_add_tcPr()
    va = OxmlElement("w:vAlign")
    va.set(qn("w:val"), "center")
    tcPr.append(va)
    trHeight = tbl.rows[0]._tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(int(height_cm * 567)))
    h.set(qn("w:hRule"), "atLeast")
    trHeight.append(h)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("🖼  " + label)
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = GREY
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("[ insert screenshot here ]")
    r2.font.size = Pt(8)
    r2.font.italic = True
    r2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    add_para(caption, size=9, italic=True, color=GREY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10, space_before=2)


def add_table(headers, rows, widths=None, header_shade=HEAD_SHADE,
              font_size=10, col_align=None):
    n = len(headers)
    tbl = doc.add_table(rows=1, cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    if widths is None:
        widths = [CONTENT_WIDTH / n] * n
    hdr = tbl.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].width = widths[i]
        shade_cell(hdr[i], header_shade)
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(htext)
        r.font.bold = True
        r.font.size = Pt(font_size)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].width = widths[i]
            p = cells[i].paragraphs[0]
            if col_align and col_align[i]:
                p.alignment = col_align[i]
            r = p.add_run(str(val))
            r.font.size = Pt(font_size)
    add_para("", space_after=6)
    return tbl


def add_field(paragraph, field_code):
    """Insert a Word field (e.g. PAGE, TOC) into a paragraph."""
    run = paragraph.add_run()
    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fldSep = OxmlElement("w:fldChar")
    fldSep.set(qn("w:fldCharType"), "separate")
    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")
    run._r.append(fldBegin)
    run._r.append(instr)
    run._r.append(fldSep)
    run._r.append(fldEnd)


def hrule_after_heading():
    add_para("", space_after=2)


# ============================================================
# COVER PAGE
# ============================================================
add_para("", space_before=40)
add_para("MOBILE APPLICATION DEVELOPMENT", size=12, bold=True, color=GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para("Project Report", size=12, italic=True, color=GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)

p = add_para(align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
r = p.add_run("My Wallet")
r.font.size = Pt(40)
r.font.bold = True
r.font.color.rgb = AMBER

add_para("Personal Expense Tracker", size=20, bold=True, color=DARK,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para("A Cross-Platform CRUD Mobile Application Built with Flutter",
         size=12, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_after=30)

add_placeholder("App Home Screen (Ledger)",
                "Figure 1: The My Wallet home screen.",
                width_cm=7.5, height_cm=6.0)

add_para("", space_after=20)
# Group box
add_para("Submitted by Group [Group Name / Number]", size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
add_table(
    ["No.", "Member Name", "Student ID", "Primary Role"],
    [
        ["1", "Member 1", "[Student ID]", "Project Lead / State Management"],
        ["2", "Member 2", "[Student ID]", "UI/UX — Ledger & Detail"],
        ["3", "Member 3", "[Student ID]", "Charts & Discover"],
        ["4", "Member 4", "[Student ID]", "Persistence & Testing"],
        ["5", "Member 5", "[Student ID]", "Documentation & Presentation"],
    ],
    widths=[Cm(1.2), Cm(5.0), Cm(3.3), Cm(6.5)],
    font_size=10,
)
add_para("", space_after=10)
add_para("Course: Mobile Application Development        Semester: [Semester]",
         size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para("Lecturer: [Lecturer Name]        Date: June 2026",
         size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ============================================================
# TABLE OF CONTENTS
# ============================================================
add_para("Table of Contents", size=18, bold=True, color=DARK, space_after=10)
toc_p = doc.add_paragraph()
add_field(toc_p, 'TOC \\o "1-3" \\h \\z \\u')
add_para("(In Microsoft Word, right-click the table above and choose "
         "“Update Field” to generate page numbers.)",
         size=9, italic=True, color=GREY, space_before=6)
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ============================================================
# 1. INTRODUCTION
# ============================================================
doc.add_heading("1. Introduction", level=1)

doc.add_heading("1.1 Project Background", level=2)
add_para(
    "Managing personal finances is a daily challenge for students and young "
    "working adults. Small, frequent purchases — a coffee here, a bus fare there — "
    "are easy to forget and quickly add up, making it difficult to understand where "
    "money actually goes each month. While many commercial finance apps exist, they "
    "are often cluttered with features, require account registration, or push the "
    "user's data to the cloud."
)
add_para(
    "My Wallet is a lightweight personal expense tracker that solves this problem. "
    "It lets a user record income and expenses in seconds, review them grouped by day "
    "and month, visualise spending trends, and stay within a monthly budget — all "
    "while keeping every record stored privately on the device. The application was "
    "developed with Flutter, Google's cross-platform UI toolkit, allowing a single "
    "codebase to target Android, iOS, web and desktop."
)

doc.add_heading("1.2 Problem Statement", level=2)
add_bullets([
    "Users lack a fast, friction-free way to log everyday transactions.",
    "Existing apps are often too complex, ad-heavy, or require online accounts.",
    "Users want to see spending patterns and control a monthly budget at a glance.",
    "Privacy-conscious users prefer their financial data to stay on their own device.",
])

doc.add_heading("1.3 Objectives", level=2)
add_numbers([
    "To design and develop a mobile application implementing full CRUD "
    "(Create, Read, Update, Delete) functionality.",
    "To apply local state management using the Provider pattern.",
    "To build a clean, responsive, and intuitive user interface using Flutter's "
    "Material 3 widgets.",
    "To persist data locally so that records survive app restarts.",
    "To practise team collaboration and version control using Git and GitHub.",
])

doc.add_heading("1.4 Scope", level=2)
add_para(
    "The project covers a fully functional expense-tracking app with transaction "
    "management, monthly browsing, charts, budget tracking, and a profile/statistics "
    "page. The following are intentionally out of scope for this assignment: "
    "multi-user accounts, cloud synchronisation, bank integration, and online "
    "authentication. These are discussed as future enhancements in Section 10."
)

doc.add_heading("1.5 Target Users", level=2)
add_bullets([
    ("Students — ", "track allowance and daily spending on a tight budget."),
    ("Young working adults — ", "monitor monthly income, expenses and savings."),
    ("Privacy-minded users — ", "anyone who wants offline, on-device finance tracking."),
])

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ============================================================
# 2. APPLICATION OVERVIEW
# ============================================================
doc.add_heading("2. Application Overview", level=1)

doc.add_heading("2.1 Theme and Concept", level=2)
add_para(
    "From the list of suggested themes, the team selected the Personal Expense "
    "Tracker. The concept is a digital ledger: each entry is a single transaction "
    "with an amount, a category, a date, a type (income or expense) and an optional "
    "note. The interface borrows the familiar “daily ledger” metaphor used "
    "by popular bookkeeping apps, where transactions are grouped under date headers "
    "and a yellow header summarises the month's income and expense."
)

doc.add_heading("2.2 Core Features (CRUD)", level=2)
add_table(
    ["CRUD Operation", "Feature in My Wallet"],
    [
        ["Create", "Add a transaction through a custom numeric keypad, category "
         "grid, date picker and note field."],
        ["Read", "Browse transactions grouped by day in the Ledger; open any "
         "record in a full Detail view."],
        ["Update", "Edit an existing transaction — the Add screen is reused in "
         "“edit” mode, pre-filled with the record's data."],
        ["Delete", "Remove a single record with a confirmation dialog, or clear "
         "all records from the Profile page."],
    ],
    widths=[Cm(3.5), Cm(12.5)],
)

doc.add_heading("2.3 Additional Features", level=2)
add_bullets([
    ("Charts — ", "a custom-painted line chart (Week / Month / Year) plus a "
     "category breakdown with percentages and progress bars."),
    ("Monthly budget — ", "set a spending limit and track the remaining balance "
     "with a circular progress indicator."),
    ("Statistics — ", "totals for income, expense and net balance, plus the number "
     "of records and active days."),
    ("Month navigation — ", "step between months or jump to any month via a picker."),
    ("Local persistence — ", "all data is stored on-device with shared_preferences."),
])

doc.add_heading("2.4 Technology Stack", level=2)
add_table(
    ["Layer", "Technology", "Purpose"],
    [
        ["Language", "Dart (SDK ^3.11)", "Application programming language"],
        ["Framework", "Flutter (Material 3)", "Cross-platform UI toolkit"],
        ["State", "provider ^6.1", "ChangeNotifier-based state management"],
        ["Storage", "shared_preferences ^2.3", "Local key/value persistence"],
        ["Utilities", "intl ^0.20", "Date and number formatting"],
        ["Utilities", "uuid ^4.5", "Unique transaction IDs"],
        ["Version control", "Git + GitHub", "Source control and collaboration"],
        ["Testing", "flutter_test", "Unit and widget testing"],
    ],
    widths=[Cm(3.2), Cm(5.0), Cm(7.8)],
)

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ============================================================
# 3. SYSTEM ARCHITECTURE
# ============================================================
doc.add_heading("3. System Architecture", level=1)

doc.add_heading("3.1 Architectural Overview", level=2)
add_para(
    "My Wallet follows a simple layered architecture that cleanly separates the user "
    "interface, the application state, the data model, and storage. The UI never "
    "talks to storage directly; instead it reads from and writes to a single "
    "ChangeNotifier (TransactionProvider), which is the single source of truth. This "
    "keeps widgets stateless where possible and makes the logic easy to test."
)
add_code(
    "  +-------------------------------------------------+\n"
    "  |        UI Layer  (Screens & Widgets)            |\n"
    "  |  Ledger | Charts | Discover | Profile | AddEdit |\n"
    "  +------------------------+------------------------+\n"
    "            context.watch() |  context.read()\n"
    "                            v\n"
    "  +-------------------------------------------------+\n"
    "  |   State Layer  (TransactionProvider)            |\n"
    "  |   add() update() delete() totals budget         |\n"
    "  +------------------------+------------------------+\n"
    "            toJson()        |  fromJson()\n"
    "                            v\n"
    "  +-------------------------------------------------+\n"
    "  |   Model Layer (Transaction, Category, Type)     |\n"
    "  +------------------------+------------------------+\n"
    "                            v\n"
    "  +-------------------------------------------------+\n"
    "  |   Storage Layer  (SharedPreferences)            |\n"
    "  +-------------------------------------------------+",
    "Figure 2: Layered architecture of My Wallet.",
)

doc.add_heading("3.2 Project Folder Structure", level=2)
add_code(
    "lib/\n"
    "  main.dart                     # App entry, theme, Provider injection\n"
    "  models/\n"
    "    transaction.dart            # Transaction model + enums\n"
    "  providers/\n"
    "    transaction_provider.dart   # CRUD logic, totals, budget, persistence\n"
    "  screens/\n"
    "    main_screen.dart            # Bottom-navigation shell\n"
    "    ledger_screen.dart          # READ  - list grouped by day\n"
    "    add_record_screen.dart      # CREATE & UPDATE - keypad + categories\n"
    "    detail_screen.dart          # READ one + entry to UPDATE / DELETE\n"
    "    charts_screen.dart          # Line chart + category breakdown\n"
    "    discover_screen.dart        # Monthly bill + budget tracking\n"
    "    profile_screen.dart         # Stats, settings, clear-all\n"
    "  widgets/\n"
    "    transaction_tile.dart       # Reusable list row\n"
    "test/\n"
    "    transaction_model_test.dart\n"
    "    transaction_provider_test.dart\n"
    "    widget_test.dart",
    "Figure 3: Source-code folder structure.",
)

doc.add_heading("3.3 Data Model", level=2)
add_para(
    "The core data structure is the Transaction class. Each transaction stores an "
    "id, title, amount, category, date, type and note. Two enums — TransactionType "
    "(income / expense) and Category (food, transport, shopping, etc.) — keep the "
    "data type-safe. The model includes toMap/fromMap and toJson/fromJson methods so "
    "it can be serialised to and from the local store."
)
add_code(
    "class Transaction {\n"
    "  final String id;\n"
    "  final String title;\n"
    "  final double amount;\n"
    "  final Category category;\n"
    "  final DateTime date;\n"
    "  final TransactionType type;\n"
    "  final String note;\n"
    "\n"
    "  Map<String, dynamic> toMap() => {\n"
    "        'id': id, 'title': title, 'amount': amount,\n"
    "        'category': category.index,\n"
    "        'date': date.millisecondsSinceEpoch,\n"
    "        'type': type.index, 'note': note,\n"
    "      };\n"
    "\n"
    "  factory Transaction.fromMap(Map<String, dynamic> map) => Transaction(\n"
    "        id: map['id'], title: map['title'], amount: map['amount'],\n"
    "        category: Category.values[map['category']],\n"
    "        date: DateTime.fromMillisecondsSinceEpoch(map['date']),\n"
    "        type: TransactionType.values[map['type']],\n"
    "        note: map['note'] ?? '',\n"
    "      );\n"
    "}",
    "Listing 1: The Transaction model (lib/models/transaction.dart).",
)

doc.add_heading("3.4 State Management", level=2)
add_para(
    "State is managed with the Provider package. TransactionProvider extends "
    "ChangeNotifier and holds the in-memory list of transactions plus the monthly "
    "budget. It exposes computed getters (totalIncome, totalExpense, balance) and "
    "the four CRUD methods. Whenever the data changes, notifyListeners() rebuilds any "
    "widget that is watching the provider. The provider is injected once at the top of "
    "the widget tree in main.dart using ChangeNotifierProvider."
)
add_code(
    "ChangeNotifierProvider(\n"
    "  create: (_) => TransactionProvider(),\n"
    "  child: MaterialApp(home: const MainScreen()),\n"
    ");",
    "Listing 2: Injecting the provider in main.dart.",
)

doc.add_heading("3.5 Data Persistence", level=2)
add_para(
    "Persistence is handled by shared_preferences. The full list of transactions is "
    "stored as a list of JSON strings under a single key. On startup the provider "
    "loads and decodes this list; after every add, update or delete it re-encodes and "
    "saves. This approach is simple, dependency-light and perfectly adequate for the "
    "data volumes of a personal tracker."
)
add_code(
    "Future<void> _save() async {\n"
    "  final prefs = await SharedPreferences.getInstance();\n"
    "  await prefs.setStringList(\n"
    "    'transactions',\n"
    "    _transactions.map((t) => t.toJson()).toList(),\n"
    "  );\n"
    "}",
    "Listing 3: Saving transactions to local storage.",
)

doc.add_heading("3.6 Navigation Structure", level=2)
add_para(
    "The app uses a custom bottom navigation bar with four primary tabs — Ledger, "
    "Charts, Discover and Profile — plus a central circular “Record” button "
    "that opens the Add screen as a full-screen modal. An IndexedStack preserves the "
    "scroll position and state of each tab when the user switches between them. The "
    "Detail screen is pushed on top of the navigation stack with MaterialPageRoute."
)

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ============================================================
# 4. UI/UX DESIGN
# ============================================================
doc.add_heading("4. UI / UX Design", level=1)

doc.add_heading("4.1 Design Principles", level=2)
add_bullets([
    ("Speed first — ", "the most common action, recording a transaction, is one tap "
     "away via the central button and a built-in numeric keypad."),
    ("Familiarity — ", "a daily-ledger layout and Material 3 components match users' "
     "existing mental models."),
    ("Clarity — ", "income is green, expense is red, and the amber brand colour "
     "highlights primary actions."),
    ("Consistency — ", "rounded cards, the same header style, and reusable widgets "
     "across every screen."),
])

doc.add_heading("4.2 Colour Scheme and Theming", level=2)
add_para(
    "The app is themed around an amber seed colour (#FFC107) using Material 3's "
    "ColorScheme.fromSeed, with a light grey scaffold background (#F5F5F5). Income and "
    "expense use semantic green and red. The theme is defined once in main.dart and "
    "applied app-wide."
)
add_table(
    ["Colour", "Hex", "Usage"],
    [
        ["Amber (brand)", "#FFC107", "Headers, primary buttons, selected states"],
        ["Light grey", "#F5F5F5", "Screen background"],
        ["Green", "#43A047", "Income amounts, positive balance"],
        ["Red", "#E53935", "Expense amounts, delete actions"],
        ["Dark grey", "#222222", "Primary text"],
    ],
    widths=[Cm(3.5), Cm(2.8), Cm(9.7)],
)

doc.add_heading("4.3 Screen-by-Screen Walkthrough", level=2)

doc.add_heading("4.3.1 Ledger (Home)", level=3)
add_para(
    "The Ledger is the landing screen. A yellow header shows a month selector and the "
    "selected month's total income and expense. Below it, transactions are grouped "
    "under date headers, each group rendered as a white rounded card. Tapping a row "
    "opens its Detail screen.")
add_placeholder("Ledger Screen", "Figure 4: Transactions grouped by day.",
                width_cm=7.0, height_cm=5.5)

doc.add_heading("4.3.2 Add / Edit Record", level=3)
add_para(
    "Opened from the central Record button (or the edit icon on a Detail screen), this "
    "screen has an Expense/Income tab bar, a live amount display, a category grid, a "
    "date picker, a note field, and a custom numeric keypad with a Save key. In edit "
    "mode every field is pre-filled with the existing record.")
add_placeholder("Add / Edit Record Screen",
                "Figure 5: Recording a transaction with the keypad.",
                width_cm=7.0, height_cm=5.5)

doc.add_heading("4.3.3 Detail", level=3)
add_para(
    "The Detail screen presents a single transaction in full: a large amount card, the "
    "type, category, date, time and note, plus Edit and Delete actions in the app bar "
    "and a prominent Delete button.")
add_placeholder("Detail Screen", "Figure 6: Full record with edit/delete actions.",
                width_cm=7.0, height_cm=5.5)

doc.add_heading("4.3.4 Charts", level=3)
add_para(
    "Charts visualises spending. A Week/Month/Year toggle drives a custom-painted line "
    "chart (drawn with CustomPainter), and a category breakdown lists each category's "
    "total, percentage and a progress bar.")
add_placeholder("Charts Screen", "Figure 7: Trend line and category breakdown.",
                width_cm=7.0, height_cm=5.5)

doc.add_heading("4.3.5 Discover and Profile", level=3)
add_para(
    "Discover shows the monthly bill and a budget card with a circular progress "
    "indicator. Profile shows lifetime statistics, a financial summary, settings, and "
    "a “Clear All Records” option.")
add_placeholder("Discover & Profile Screens",
                "Figure 8: Budget tracking and statistics.",
                width_cm=7.0, height_cm=5.5)

doc.add_heading("4.4 Responsive Design", level=2)
add_para(
    "The layout adapts to different screen sizes through Flutter's flexible widgets. "
    "Expanded and Flexible distribute space in rows; GridView with a fixed cross-axis "
    "count reflows category icons; SafeArea avoids notches and system bars; and "
    "scrollable views (ListView, SingleChildScrollView) prevent overflow on small "
    "devices. Because no widths are hard-coded to a single device size, the same UI "
    "renders correctly on phones of different resolutions as well as on tablets."
)

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ============================================================
# 5. CRUD IMPLEMENTATION
# ============================================================
doc.add_heading("5. CRUD Implementation", level=1)
add_para(
    "All four CRUD operations are implemented inside TransactionProvider and are "
    "triggered from the UI. Each operation mutates the in-memory list, persists the "
    "change, and calls notifyListeners() so the UI refreshes automatically."
)

doc.add_heading("5.1 Create", level=2)
add_para(
    "When the user taps Save on the Add screen, a new Transaction is created with a "
    "UUID and passed to provider.add().")
add_code(
    "Future<void> add(Transaction t) async {\n"
    "  _transactions.add(t);\n"
    "  await _save();\n"
    "  notifyListeners();\n"
    "}",
    "Listing 4: Create operation.",
)

doc.add_heading("5.2 Read", level=2)
add_para(
    "The Ledger reads the provider with context.watch(), filters by the selected "
    "month, sorts by date, and groups by day. The Detail screen reads a single record "
    "by id via getById().")
add_code(
    "final all = provider.allTransactions\n"
    "    .where((t) => t.date.year == _month.year &&\n"
    "                  t.date.month == _month.month)\n"
    "    .toList()\n"
    "  ..sort((a, b) => b.date.compareTo(a.date));",
    "Listing 5: Reading and filtering transactions.",
)

doc.add_heading("5.3 Update", level=2)
add_para(
    "Editing reuses the Add screen. On save, provider.update() finds the record by id "
    "and replaces it.")
add_code(
    "Future<void> update(Transaction t) async {\n"
    "  final i = _transactions.indexWhere((x) => x.id == t.id);\n"
    "  if (i != -1) {\n"
    "    _transactions[i] = t;\n"
    "    await _save();\n"
    "    notifyListeners();\n"
    "  }\n"
    "}",
    "Listing 6: Update operation.",
)

doc.add_heading("5.4 Delete", level=2)
add_para(
    "Deletion is guarded by a confirmation dialog. On confirm, provider.delete() "
    "removes the record by id.")
add_code(
    "Future<void> delete(String id) async {\n"
    "  _transactions.removeWhere((t) => t.id == id);\n"
    "  await _save();\n"
    "  notifyListeners();\n"
    "}",
    "Listing 7: Delete operation.",
)

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ============================================================
# 6. TESTING
# ============================================================
doc.add_heading("6. Testing", level=1)

doc.add_heading("6.1 Testing Strategy", level=2)
add_para(
    "The team used a combination of automated tests and manual testing. Automated "
    "tests focus on the model and the provider — the parts that hold the business "
    "logic — while a widget smoke test confirms the app builds and renders. Manual "
    "testing covered end-to-end user flows on a device/emulator."
)

doc.add_heading("6.2 Automated Tests", level=2)
add_para("Three test files were written under the test/ folder:")
add_bullets([
    ("transaction_model_test.dart — ", "serialisation round-trips (toMap/fromMap, "
     "toJson/fromJson) and copyWith behaviour."),
    ("transaction_provider_test.dart — ", "add, update, delete, clearAll, setBudget, "
     "and the income/expense/balance getters."),
    ("widget_test.dart — ", "a smoke test that builds the app and verifies the bottom "
     "navigation renders."),
])
add_para("All 11 automated tests pass, as summarised below.", space_before=4)
add_table(
    ["#", "Test Case", "Type", "Result"],
    [
        ["1", "toMap / fromMap round trip", "Unit", "Pass"],
        ["2", "toJson / fromJson round trip", "Unit", "Pass"],
        ["3", "copyWith overrides only given fields", "Unit", "Pass"],
        ["4", "Category exposes label and icon", "Unit", "Pass"],
        ["5", "Provider starts empty", "Unit", "Pass"],
        ["6", "add() updates list and totals", "Unit", "Pass"],
        ["7", "update() replaces matching record", "Unit", "Pass"],
        ["8", "delete() removes matching record", "Unit", "Pass"],
        ["9", "clearAll() empties the ledger", "Unit", "Pass"],
        ["10", "setBudget() stores the budget", "Unit", "Pass"],
        ["11", "App builds and shows navigation", "Widget", "Pass"],
    ],
    widths=[Cm(1.0), Cm(8.5), Cm(3.0), Cm(3.5)],
    col_align=[WD_ALIGN_PARAGRAPH.CENTER, None,
               WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
)
add_code(
    "$ flutter test\n"
    "00:00 +11: All tests passed!",
    "Listing 8: Test run output.",
)

doc.add_heading("6.3 Manual Testing", level=2)
add_table(
    ["#", "Scenario", "Expected Result", "Result"],
    [
        ["M1", "Add an expense via the keypad", "Record appears in today's group; "
         "month expense increases", "Pass"],
        ["M2", "Add an income", "Record appears; month income and balance increase",
         "Pass"],
        ["M3", "Edit a record's amount", "Detail and Ledger show the new amount",
         "Pass"],
        ["M4", "Delete a record", "Confirmation shown; record removed after confirm",
         "Pass"],
        ["M5", "Switch months", "Ledger shows only the selected month's records",
         "Pass"],
        ["M6", "Set a monthly budget", "Budget card shows progress and remaining",
         "Pass"],
        ["M7", "Restart the app", "All records persist (loaded from storage)",
         "Pass"],
        ["M8", "Clear all records", "Ledger is emptied after confirmation", "Pass"],
    ],
    widths=[Cm(1.0), Cm(4.8), Cm(7.7), Cm(2.5)],
    font_size=9.5,
    col_align=[WD_ALIGN_PARAGRAPH.CENTER, None, None, WD_ALIGN_PARAGRAPH.CENTER],
)

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ============================================================
# 7. VERSION CONTROL
# ============================================================
doc.add_heading("7. Version Control", level=1)
add_para(
    "The project is managed with Git and hosted on GitHub. Version control let the "
    "team work in parallel, track the history of every change, and integrate work "
    "safely."
)
doc.add_heading("7.1 Workflow", level=2)
add_bullets([
    "A shared GitHub repository hosts the codebase.",
    "Each member works on their own area (screens, provider, tests, docs) and commits "
    "small, focused changes with descriptive messages.",
    "A .gitignore excludes build artefacts (build/, .dart_tool/) so only source is "
    "tracked.",
    "The README documents the project, architecture, and how to run it.",
])
doc.add_heading("7.2 Commit Convention", level=2)
add_para("Commit messages follow a short, consistent style, for example:")
add_code(
    "feat: add monthly budget card to Discover screen\n"
    "fix: prevent negative amounts on the keypad\n"
    "test: add provider CRUD unit tests\n"
    "docs: write project README",
    "Listing 9: Example commit messages.",
)

doc.add_heading("8. Group Contributions", level=1)
add_para(
    "Every member contributed to design, development, testing and documentation. The "
    "table below summarises each member's primary responsibilities and estimated "
    "contribution."
)
add_table(
    ["Member", "Primary Responsibilities", "Contribution"],
    [
        ["Member 1", "Project coordination; data model and TransactionProvider "
         "(state management, CRUD, persistence).", "20%"],
        ["Member 2", "UI/UX design; Ledger and Detail screens; reusable "
         "TransactionTile widget.", "20%"],
        ["Member 3", "Add/Edit Record screen (keypad, category grid) and the Charts "
         "screen (custom line chart).", "20%"],
        ["Member 4", "Discover (budget) and Profile screens; persistence wiring; "
         "automated tests.", "20%"],
        ["Member 5", "Documentation, README, project report and presentation "
         "slides; manual testing.", "20%"],
    ],
    widths=[Cm(2.6), Cm(10.4), Cm(3.0)],
    col_align=[None, None, WD_ALIGN_PARAGRAPH.CENTER],
)
add_para("Note: Replace the placeholder names and adjust the contribution split to "
         "reflect your group's actual work.", size=9, italic=True, color=GREY)

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ============================================================
# 9. CHALLENGES
# ============================================================
doc.add_heading("9. Challenges and Solutions", level=1)
add_table(
    ["Challenge", "Solution"],
    [
        ["Keeping the UI in sync with data changes across multiple screens.",
         "Used a single ChangeNotifier (Provider) as the source of truth; "
         "notifyListeners() rebuilds all watching widgets automatically."],
        ["Persisting structured data with a lightweight dependency.",
         "Serialised each Transaction to JSON and stored the list with "
         "shared_preferences, avoiding a full database."],
        ["Drawing a trend chart without a heavy charting library.",
         "Implemented a custom CustomPainter to draw the line chart, average line "
         "and data points."],
        ["Preventing invalid input on the amount keypad.",
         "Added input rules: a single decimal point, at most two decimals, and a "
         "maximum length; Save rejects a zero amount."],
        ["Avoiding losing each tab's scroll state when switching tabs.",
         "Wrapped the pages in an IndexedStack so all four tabs stay alive."],
    ],
    widths=[Cm(7.0), Cm(9.0)],
    font_size=10,
)

doc.add_heading("10. Conclusion and Future Enhancements", level=1)
doc.add_heading("10.1 Conclusion", level=2)
add_para(
    "My Wallet successfully meets all of the assignment's functional requirements. It "
    "delivers complete CRUD functionality, a clean and responsive Material 3 "
    "interface, Provider-based local state management, and on-device persistence. "
    "Building the app gave the team hands-on experience with Flutter, the Provider "
    "pattern, automated testing, and collaborative development with Git and GitHub."
)
doc.add_heading("10.2 Future Enhancements", level=2)
add_bullets([
    "Cloud sync and multi-device support via a backend or Firebase.",
    "Optional user accounts with secure authentication.",
    "Export to CSV/PDF and receipt photo attachments.",
    "Recurring transactions and bill reminders (notifications).",
    "Dark mode and multiple-currency support.",
    "Migrating storage from shared_preferences to a SQLite database (sqflite/drift) "
    "for larger data volumes and richer queries.",
])

doc.add_heading("11. References", level=1)
refs = [
    "Flutter Documentation. https://docs.flutter.dev",
    "Dart Language Tour. https://dart.dev/guides",
    "Provider package. https://pub.dev/packages/provider",
    "shared_preferences package. https://pub.dev/packages/shared_preferences",
    "intl package. https://pub.dev/packages/intl",
    "Material Design 3. https://m3.material.io",
]
for r in refs:
    p = doc.add_paragraph(style="List Number")
    p.add_run(r)
    p.paragraph_format.space_after = Pt(3)

# Appendix
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
doc.add_heading("Appendix A: Application Screenshots", level=1)
add_para("Insert the captured screenshots of the running application in the boxes "
         "below.", italic=True, color=GREY)
screens = [
    ("Ledger (Home)", "Transactions grouped by day."),
    ("Add / Edit Record", "Numeric keypad and category grid."),
    ("Detail", "Single record with edit and delete."),
    ("Charts", "Trend line and category breakdown."),
    ("Discover", "Monthly bill and budget tracking."),
    ("Profile", "Statistics, summary and settings."),
]
# 2-column grid of placeholders
grid = doc.add_table(rows=0, cols=2)
grid.alignment = WD_TABLE_ALIGNMENT.CENTER
for i in range(0, len(screens), 2):
    cells = grid.add_row().cells
    for j in range(2):
        if i + j < len(screens):
            name, cap = screens[i + j]
            c = cells[j]
            c.width = Cm(8.0)
            shade_cell(c, "FAFAFA")
            set_cell_borders(c, color="BBBBBB", sz="8")
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = p.add_run("🖼  " + name)
            rr.font.bold = True
            rr.font.color.rgb = GREY
            p2 = c.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2 = p2.add_run("[ insert screenshot ]")
            r2.font.size = Pt(8)
            r2.font.italic = True
            r2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
            p3 = c.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r3 = p3.add_run(cap)
            r3.font.size = Pt(9)
            r3.font.italic = True
            r3.font.color.rgb = GREY
            # row height
            tr = grid.rows[-1]._tr.get_or_add_trPr()
            h = OxmlElement("w:trHeight")
            h.set(qn("w:val"), str(int(5.0 * 567)))
            h.set(qn("w:hRule"), "atLeast")
            tr.append(h)

# ============================================================
# FOOTER with page numbers
# ============================================================
footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.add_run("My Wallet — Personal Expense Tracker   |   Page ")
add_field(fp, "PAGE")
fp.add_run(" of ")
add_field(fp, "NUMPAGES")
for run in fp.runs:
    run.font.size = Pt(8)
    run.font.color.rgb = GREY

out = r"D:\des3113\docs\My_Wallet_Project_Report.docx"
doc.save(out)
print("Saved:", out)
print("Paragraphs:", len(doc.paragraphs), "Tables:", len(doc.tables))
